import os
import io
import csv
import uuid
import re
import shutil
import subprocess
import sys
from datetime import datetime, date
from functools import wraps
from pathlib import Path

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, jsonify, flash, send_from_directory, abort, Response
)
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash

import pdfplumber
import openpyxl
import xlrd
from docx import Document

from database import get_db, init_db

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
TEMPLATES_DIR = os.path.join(UPLOAD_DIR, "templates")
GENERATED_DIR = os.path.join(UPLOAD_DIR, "generated")
ALLOWED_EXT = {"png", "jpg", "jpeg", "pdf"}
MAX_CONTENT_LENGTH = 8 * 1024 * 1024  # 8 MB

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "cambia-esta-clave-en-produccion")
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)


@app.template_filter("basename")
def basename_filter(s):
    return os.path.basename(s) if s else ""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT


def parse_date(s):
    return datetime.strptime(s, "%Y-%m-%d").date()


def enviar_notificacion(empleado, tipo_registro, estado_anterior, nuevo_estado, folio):
    # Convert Row or dict to dictionary safely
    emp_dict = dict(empleado) if empleado else {}
    correo = emp_dict.get("correo")
    telefono = emp_dict.get("telefono")
    nombre = emp_dict.get("nombre")
    
    if not correo and not telefono:
        print(f"[Notificación] Sin medios de contacto para {nombre}. Notificación no enviada.")
        return False

    mensaje = ""
    if nuevo_estado == "creado":
        if tipo_registro == "incapacidad":
            mensaje = f"Estimado/a {nombre}, su solicitud de incapacidad (Folio #{folio}) ha sido recibida exitosamente en el sistema."
        else:
            mensaje = f"Estimado/a {nombre}, se ha registrado su solicitud de vacaciones (Folio #{folio}) y se encuentra en estado 'Pendiente' para aprobación."
    elif nuevo_estado == "aprobado":
        mensaje = f"Estimado/a {nombre}, su solicitud de {tipo_registro} (Folio #{folio}) ha sido APROBADA."
    elif nuevo_estado == "rechazado":
        mensaje = f"Estimado/a {nombre}, su solicitud de {tipo_registro} (Folio #{folio}) ha sido RECHAZADA."

    print(f"\n======== NOTIFICACIÓN DISPACHADA ========")
    if correo:
        print(f"PARA (Email): {correo}")
    if telefono:
        print(f"PARA (Teléfono/SMS): {telefono}")
    print(f"MENSAJE: {mensaje}")
    print(f"=========================================\n")
    return True


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("admin_id"):
            return redirect(url_for("admin_login", next=request.path))
        return view(*args, **kwargs)
    return wrapped


# ---------------------------------------------------------------------------
# Rutas publicas / empleado
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/empleado")
def empleado_form():
    return render_template("empleado.html", today=date.today().isoformat())


@app.route("/api/empleado/<codigo>")
def api_empleado(codigo):
    codigo = codigo.strip().zfill(4) if codigo.strip().isdigit() else codigo.strip()
    db = get_db()
    row = db.execute(
        "SELECT codigo, nombre, cargo, fecha_ingreso, activo, dui, correo, telefono FROM empleados WHERE codigo = ?",
        (codigo,),
    ).fetchone()
    db.close()
    if not row:
        return jsonify({"ok": False, "error": "Codigo no encontrado"}), 404
    if not row["activo"]:
        return jsonify({"ok": False, "error": "Empleado inactivo"}), 404
    return jsonify({
        "ok": True,
        "codigo": row["codigo"],
        "nombre": row["nombre"],
        "cargo": row["cargo"],
        "fecha_ingreso": row["fecha_ingreso"],
        "dui": row["dui"],
        "correo": row["correo"],
        "telefono": row["telefono"]
    })


@app.route("/api/empleado/dui/<dui>")
def api_empleado_by_dui(dui):
    dui = dui.strip()
    db = get_db()
    row = db.execute(
        "SELECT codigo, nombre, cargo, fecha_ingreso, activo, dui, correo, telefono FROM empleados WHERE dui = ?",
        (dui,),
    ).fetchone()
    db.close()
    if not row:
        return jsonify({"ok": False, "error": "DUI no encontrado"}), 404
    if not row["activo"]:
        return jsonify({"ok": False, "error": "Empleado inactivo"}), 404
    return jsonify({
        "ok": True,
        "codigo": row["codigo"],
        "nombre": row["nombre"],
        "cargo": row["cargo"],
        "fecha_ingreso": row["fecha_ingreso"],
        "dui": row["dui"],
        "correo": row["correo"],
        "telefono": row["telefono"]
    })


@app.route("/api/registro", methods=["POST"])
def api_registro():
    dui = request.form.get("dui", "").strip()
    nombre = request.form.get("nombre", "").strip()
    cargo = request.form.get("cargo", "").strip()
    correo = request.form.get("correo", "").strip()
    telefono = request.form.get("telefono", "").strip()
    tipo = request.form.get("tipo", "").strip()
    fecha_inicio_s = request.form.get("fecha_inicio", "").strip()
    fecha_fin_s = request.form.get("fecha_fin", "").strip()
    observaciones = request.form.get("observaciones", "").strip()

    if not dui:
        return jsonify({"ok": False, "error": "El DUI es obligatorio"}), 400

    if tipo not in ("vacacion", "incapacidad"):
        return jsonify({"ok": False, "error": "Tipo invalido"}), 400

    db = get_db()
    
    # Try to find employee by DUI
    emp = db.execute("SELECT * FROM empleados WHERE dui = ?", (dui,)).fetchone()
    
    if not emp:
        # If employee doesn't exist, we must create a new one!
        if not nombre or not cargo:
            db.close()
            return jsonify({"ok": False, "error": "Para un nuevo registro, el nombre y el puesto son obligatorios"}), 400
        
        # Generate next codigo
        row = db.execute("SELECT codigo FROM empleados ORDER BY CAST(codigo AS INTEGER) DESC LIMIT 1").fetchone()
        if row:
            try:
                next_val = int(row["codigo"]) + 1
                codigo = str(next_val).zfill(4)
            except ValueError:
                codigo = "2000"
        else:
            codigo = "2000"
            
        # Insert new employee
        db.execute(
            "INSERT INTO empleados (codigo, nombre, cargo, fecha_ingreso, salario_mensual, activo, dui, correo, telefono) VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?)",
            (codigo, nombre, cargo, date.today().isoformat(), 0.0, dui, correo, telefono)
        )
        db.commit()
        
        # Refetch the new employee to be safe
        emp = db.execute("SELECT * FROM empleados WHERE codigo = ?", (codigo,)).fetchone()
    else:
        codigo = emp["codigo"]
        updated = False
        new_correo = emp["correo"]
        new_telefono = emp["telefono"]
        if correo and correo != emp["correo"]:
            new_correo = correo
            updated = True
        if telefono and telefono != emp["telefono"]:
            new_telefono = telefono
            updated = True
        if updated:
            db.execute("UPDATE empleados SET correo = ?, telefono = ? WHERE codigo = ?", (new_correo, new_telefono, codigo))
            db.commit()
            emp = db.execute("SELECT * FROM empleados WHERE codigo = ?", (codigo,)).fetchone()

    try:
        fi = parse_date(fecha_inicio_s)
        ff = parse_date(fecha_fin_s)
    except ValueError:
        db.close()
        return jsonify({"ok": False, "error": "Fechas invalidas"}), 400

    if ff < fi:
        db.close()
        return jsonify({"ok": False, "error": "La fecha final no puede ser anterior a la fecha de inicio"}), 400

    dias = (ff - fi).days + 1

    foto_path = None
    if tipo == "incapacidad":
        file = request.files.get("foto")
        if not file or file.filename == "":
            db.close()
            return jsonify({"ok": False, "error": "Debes adjuntar la foto/constancia de incapacidad"}), 400
        if not allowed_file(file.filename):
            db.close()
            return jsonify({"ok": False, "error": "Formato de archivo no permitido (usa jpg, png o pdf)"}), 400
        ext = file.filename.rsplit(".", 1)[1].lower()
        safe_name = secure_filename(f"{codigo}_{fi.isoformat()}_{uuid.uuid4().hex[:8]}.{ext}")
        file.save(os.path.join(UPLOAD_DIR, safe_name))
        foto_path = safe_name

    cur = db.execute(
        """INSERT INTO registros (codigo_empleado, tipo, fecha_inicio, fecha_fin, dias, foto_path, observaciones, fecha_registro)
           VALUES (?,?,?,?,?,?,?,?)""",
        (codigo, tipo, fi.isoformat(), ff.isoformat(), dias, foto_path, observaciones,
         datetime.now().isoformat(timespec="seconds")),
    )
    db.commit()
    new_id = cur.lastrowid
    
    # Send simulated notification
    enviar_notificacion(emp, tipo, None, "creado", new_id)
    
    db.close()

    msg = "Incapacidad recibida exitosamente." if tipo == "incapacidad" else ""
    return jsonify({"ok": True, "folio": new_id, "dias": dias, "nombre": emp["nombre"], "msg": msg})


@app.route("/registro/confirmacion/<int:folio>")
def confirmacion(folio):
    db = get_db()
    row = db.execute(
        """SELECT r.*, e.nombre, e.cargo FROM registros r
           JOIN empleados e ON e.codigo = r.codigo_empleado
           WHERE r.id = ?""",
        (folio,),
    ).fetchone()
    db.close()
    if not row:
        abort(404)
    return render_template("confirmacion.html", r=row)


# ---------------------------------------------------------------------------
# Admin
# ---------------------------------------------------------------------------

@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        db = get_db()
        user = db.execute("SELECT * FROM admin_users WHERE username = ?", (username,)).fetchone()
        db.close()
        if user and check_password_hash(user["password_hash"], password):
            session["admin_id"] = user["id"]
            session["admin_username"] = user["username"]
            nxt = request.args.get("next") or url_for("admin_dashboard")
            return redirect(nxt)
        flash("Usuario o clave incorrectos", "error")
    return render_template("admin_login.html")


@app.route("/admin/logout")
def admin_logout():
    session.clear()
    return redirect(url_for("admin_login"))


@app.route("/admin")
@login_required
def admin_dashboard():
    db = get_db()

    codigo = request.args.get("codigo", "").strip()
    tipo = request.args.get("tipo", "").strip()
    desde = request.args.get("desde", "").strip()
    hasta = request.args.get("hasta", "").strip()

    query = """
        SELECT r.id, r.codigo_empleado, e.nombre, e.cargo, r.tipo, r.fecha_inicio, r.fecha_fin,
               r.dias, r.foto_path, r.observaciones, r.fecha_registro, r.estado
        FROM registros r
        JOIN empleados e ON e.codigo = r.codigo_empleado
        WHERE 1=1
    """
    params = []
    if codigo:
        query += " AND (r.codigo_empleado LIKE ? OR e.dui LIKE ?)"
        params.append(f"%{codigo}%")
        params.append(f"%{codigo}%")
    if tipo in ("vacacion", "incapacidad"):
        query += " AND r.tipo = ?"
        params.append(tipo)
    if desde:
        query += " AND r.fecha_inicio >= ?"
        params.append(desde)
    if hasta:
        query += " AND r.fecha_fin <= ?"
        params.append(hasta)
    query += " ORDER BY r.fecha_registro DESC"

    registros = db.execute(query, params).fetchall()

    total_empleados = db.execute("SELECT COUNT(*) c FROM empleados WHERE activo = 1").fetchone()["c"]
    total_vacaciones = db.execute("SELECT COUNT(*) c, COALESCE(SUM(dias),0) d FROM registros WHERE tipo='vacacion'").fetchone()
    total_incapacidades = db.execute("SELECT COUNT(*) c, COALESCE(SUM(dias),0) d FROM registros WHERE tipo='incapacidad'").fetchone()

    db.close()

    return render_template(
        "admin_dashboard.html",
        registros=registros,
        total_empleados=total_empleados,
        total_vacaciones=total_vacaciones,
        total_incapacidades=total_incapacidades,
        filtros={"codigo": codigo, "tipo": tipo, "desde": desde, "hasta": hasta},
    )


@app.route("/admin/foto/<path:filename>")
@login_required
def admin_foto(filename):
    return send_from_directory(UPLOAD_DIR, filename)


@app.route("/admin/export.csv")
@login_required
def admin_export_csv():
    db = get_db()
    rows = db.execute(
        """SELECT r.id, r.codigo_empleado, e.nombre, e.cargo, r.tipo, r.fecha_inicio, r.fecha_fin,
                  r.dias, r.observaciones, r.fecha_registro, r.estado
           FROM registros r JOIN empleados e ON e.codigo = r.codigo_empleado
           ORDER BY r.fecha_registro DESC"""
    ).fetchall()
    db.close()

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Folio", "Codigo", "Nombre", "Cargo", "Tipo", "Fecha Inicio", "Fecha Fin", "Dias", "Observaciones", "Fecha Registro", "Estado"])
    for r in rows:
        writer.writerow([r["id"], r["codigo_empleado"], r["nombre"], r["cargo"], r["tipo"],
                          r["fecha_inicio"], r["fecha_fin"], r["dias"], r["observaciones"], r["fecha_registro"], r["estado"]])

    return Response(
        buf.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=reporte_dongbu.csv"},
    )


@app.route("/admin/registro/delete/<int:id>", methods=["POST"])
@login_required
def admin_delete_registro(id):
    db = get_db()
    db.execute("DELETE FROM registros WHERE id = ?", (id,))
    db.commit()
    db.close()
    flash(f"Registro con Folio #{id} eliminado correctamente.", "ok")
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/registro/approve/<int:id>", methods=["POST"])
@login_required
def admin_approve_registro(id):
    db = get_db()
    row = db.execute(
        """SELECT r.*, e.nombre, e.correo, e.telefono FROM registros r 
           JOIN empleados e ON e.codigo = r.codigo_empleado 
           WHERE r.id = ?""", (id,)
    ).fetchone()
    if row:
        db.execute("UPDATE registros SET estado = 'aprobado' WHERE id = ?", (id,))
        db.commit()
        enviar_notificacion(row, row["tipo"], "pendiente", "aprobado", id)
    db.close()
    flash(f"El Folio #{id} ha sido aprobado.", "ok")
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/registro/reject/<int:id>", methods=["POST"])
@login_required
def admin_reject_registro(id):
    db = get_db()
    row = db.execute(
        """SELECT r.*, e.nombre, e.correo, e.telefono FROM registros r 
           JOIN empleados e ON e.codigo = r.codigo_empleado 
           WHERE r.id = ?""", (id,)
    ).fetchone()
    if row:
        db.execute("UPDATE registros SET estado = 'rechazado' WHERE id = ?", (id,))
        db.commit()
        enviar_notificacion(row, row["tipo"], "pendiente", "rechazado", id)
    db.close()
    flash(f"El Folio #{id} ha sido rechazado.", "ok")
    return redirect(url_for("admin_dashboard"))


@app.route("/admin/empleados")
@login_required
def admin_empleados():
    q = request.args.get("q", "").strip()
    db = get_db()
    if q:
        rows = db.execute(
            "SELECT * FROM empleados WHERE codigo LIKE ? OR nombre LIKE ? OR cargo LIKE ? OR dui LIKE ? ORDER BY codigo",
            (f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"),
        ).fetchall()
    else:
        rows = db.execute("SELECT * FROM empleados ORDER BY codigo").fetchall()
    db.close()
    return render_template("admin_empleados.html", empleados=rows, q=q)


@app.route("/admin/empleados/nuevo", methods=["POST"])
@login_required
def admin_nuevo_empleado():
    codigo = request.form.get("codigo", "").strip()
    dui = request.form.get("dui", "").strip()
    nombre = request.form.get("nombre", "").strip()
    cargo = request.form.get("cargo", "").strip()
    fecha_ingreso = request.form.get("fecha_ingreso", "").strip()
    salario = request.form.get("salario", "").strip()
    correo = request.form.get("correo", "").strip()
    telefono = request.form.get("telefono", "").strip()

    if not dui or not nombre or not cargo:
        flash("DUI, Nombre y Puesto son obligatorios.", "error")
        return redirect(url_for("admin_empleados"))

    db = get_db()
    
    # Check duplicate DUI
    dup = db.execute("SELECT * FROM empleados WHERE dui = ?", (dui,)).fetchone()
    if dup:
        db.close()
        flash(f"Ya existe un empleado con el DUI {dui}.", "error")
        return redirect(url_for("admin_empleados"))

    if codigo:
        dup_cod = db.execute("SELECT * FROM empleados WHERE codigo = ?", (codigo,)).fetchone()
        if dup_cod:
            db.close()
            flash(f"Ya existe un empleado con el código {codigo}.", "error")
            return redirect(url_for("admin_empleados"))
    else:
        # Generate next codigo
        row = db.execute("SELECT codigo FROM empleados ORDER BY CAST(codigo AS INTEGER) DESC LIMIT 1").fetchone()
        if row:
            try:
                next_val = int(row["codigo"]) + 1
                codigo = str(next_val).zfill(4)
            except ValueError:
                codigo = "2000"
        else:
            codigo = "2000"

    try:
        sal_val = float(salario) if salario else 0.0
    except ValueError:
        sal_val = 0.0

    if not fecha_ingreso:
        fecha_ingreso = date.today().isoformat()

    isss = request.form.get("isss", "").strip()
    afp = request.form.get("afp", "").strip()
    isr = request.form.get("isr", "").strip()
    prestamo_personal = request.form.get("personal_loan", "").strip()
    prestamo_bancario = request.form.get("bank_loan", "").strip()
    fsv = request.form.get("fsv", "").strip()
    salario_letras = request.form.get("salary_words", "").strip()

    if not isss and sal_val > 0:
        isss = f"${min(sal_val, 1000.00) * 0.03:.2f}"
    if not afp and sal_val > 0:
        afp = f"${sal_val * 0.0725:.2f}"
    if not isr and sal_val > 0:
        isss_val = min(sal_val, 1000.00) * 0.03
        afp_val = sal_val * 0.0725
        taxable = max(sal_val - isss_val - afp_val, 0)
        isr_val = calculate_isr(taxable)
        isr = f"${isr_val:.2f}"
    if not salario_letras and sal_val > 0:
        salario_letras = salary_words(str(sal_val))

    db.execute(
        """INSERT INTO empleados (
            codigo, nombre, cargo, fecha_ingreso, salario_mensual, activo, dui, correo, telefono,
            isss, afp, isr, prestamo_personal, prestamo_bancario, fsv, salario_letras
           ) VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (codigo, nombre, cargo, fecha_ingreso, sal_val, dui, correo, telefono,
         isss or "$0.00", afp or "$0.00", isr or "$0.00", prestamo_personal or "$0.00",
         prestamo_bancario or "$0.00", fsv or "$0.00", salario_letras)
    )
    db.commit()
    db.close()
    flash(f"Empleado {nombre} (Código: {codigo}) agregado correctamente.", "ok")
    return redirect(url_for("admin_empleados"))


@app.route("/admin/empleados/editar/<codigo>", methods=["POST"])
@login_required
def admin_editar_empleado(codigo):
    dui = request.form.get("dui", "").strip()
    nombre = request.form.get("nombre", "").strip()
    cargo = request.form.get("cargo", "").strip()
    fecha_ingreso = request.form.get("fecha_ingreso", "").strip()
    salario = request.form.get("salario", "").strip()
    correo = request.form.get("correo", "").strip()
    telefono = request.form.get("telefono", "").strip()

    if not nombre or not cargo:
        flash("Nombre y Puesto son obligatorios.", "error")
        return redirect(url_for("admin_empleados"))

    db = get_db()

    # Check employee exists
    emp = db.execute("SELECT * FROM empleados WHERE codigo = ?", (codigo,)).fetchone()
    if not emp:
        db.close()
        flash(f"No se encontró un empleado con código {codigo}.", "error")
        return redirect(url_for("admin_empleados"))

    # Check DUI uniqueness (if changed)
    if dui and dui != emp["dui"]:
        dup = db.execute("SELECT * FROM empleados WHERE dui = ? AND codigo != ?", (dui, codigo)).fetchone()
        if dup:
            db.close()
            flash(f"Ya existe otro empleado con el DUI {dui}.", "error")
            return redirect(url_for("admin_empleados"))

    try:
        sal_val = float(salario) if salario else emp["salario_mensual"]
    except ValueError:
        sal_val = emp["salario_mensual"]

    if not fecha_ingreso:
        fecha_ingreso = emp["fecha_ingreso"]

    isss = request.form.get("isss", "").strip()
    afp = request.form.get("afp", "").strip()
    isr = request.form.get("isr", "").strip()
    prestamo_personal = request.form.get("personal_loan", "").strip()
    prestamo_bancario = request.form.get("bank_loan", "").strip()
    fsv = request.form.get("fsv", "").strip()
    salario_letras = request.form.get("salary_words", "").strip()

    if not isss and sal_val > 0:
        isss = f"${min(sal_val, 1000.00) * 0.03:.2f}"
    if not afp and sal_val > 0:
        afp = f"${sal_val * 0.0725:.2f}"
    if not isr and sal_val > 0:
        isss_val = min(sal_val, 1000.00) * 0.03
        afp_val = sal_val * 0.0725
        taxable = max(sal_val - isss_val - afp_val, 0)
        isr_val = calculate_isr(taxable)
        isr = f"${isr_val:.2f}"
    if not salario_letras and sal_val > 0:
        salario_letras = salary_words(str(sal_val))

    db.execute(
        """UPDATE empleados SET dui = ?, nombre = ?, cargo = ?, fecha_ingreso = ?,
           salario_mensual = ?, correo = ?, telefono = ?, isss = ?, afp = ?, isr = ?,
           prestamo_personal = ?, prestamo_bancario = ?, fsv = ?, salario_letras = ?
           WHERE codigo = ?""",
        (dui or emp["dui"], nombre, cargo, fecha_ingreso, sal_val, correo, telefono,
         isss or "$0.00", afp or "$0.00", isr or "$0.00", prestamo_personal or "$0.00",
         prestamo_bancario or "$0.00", fsv or "$0.00", salario_letras, codigo)
    )
    db.commit()
    db.close()
    flash(f"Empleado {nombre} (Código: {codigo}) actualizado correctamente.", "ok")
    return redirect(url_for("admin_empleados"))


@app.route("/admin/cambiar-clave", methods=["GET", "POST"])
@login_required
def admin_cambiar_clave():
    if request.method == "POST":
        actual = request.form.get("actual", "")
        nueva = request.form.get("nueva", "")
        confirmar = request.form.get("confirmar", "")
        db = get_db()
        user = db.execute("SELECT * FROM admin_users WHERE id = ?", (session["admin_id"],)).fetchone()
        if not check_password_hash(user["password_hash"], actual):
            flash("La clave actual no es correcta", "error")
        elif len(nueva) < 6:
            flash("La nueva clave debe tener al menos 6 caracteres", "error")
        elif nueva != confirmar:
            flash("Las claves no coinciden", "error")
        else:
            db.execute("UPDATE admin_users SET password_hash = ? WHERE id = ?",
                       (generate_password_hash(nueva), user["id"]))
            db.commit()
            flash("Clave actualizada correctamente", "ok")
        db.close()
    return render_template("admin_cambiar_clave.html")


# ===========================================================================
# MÓDULO DE CONSTANCIAS (Helpers y Utilidades)
# ===========================================================================

ONES = [
    "", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve", "diez",
    "once", "doce", "trece", "catorce", "quince", "dieciseis", "diecisiete", "dieciocho", "diecinueve", "veinte",
    "veintiuno", "veintidós", "veintitres", "veinticuatro", "veinticinco", "veintiseis", "veintisiete", "veintiocho", "veintinueve"
]
TENS = ["", "", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa"]
HUNDREDS = [
    "", "ciento", "doscientos", "trescientos", "cuatrocientos", "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos"
]

def number_to_words_es(n: int) -> str:
    if n == 0:
        return "cero"
    if n == 100:
        return "cien"
    if n < 30:
        return ONES[n]
    if n < 100:
        ten, one = divmod(n, 10)
        return TENS[ten] if one == 0 else f"{TENS[ten]} y {ONES[one]}"
    if n < 1000:
        hundred, rest = divmod(n, 100)
        return HUNDREDS[hundred] if rest == 0 else f"{HUNDREDS[hundred]} {number_to_words_es(rest)}"
    if n < 1_000_000:
        thousands, rest = divmod(n, 1000)
        prefix = "mil" if thousands == 1 else f"{number_to_words_es(thousands)} mil"
        return prefix if rest == 0 else f"{prefix} {number_to_words_es(rest)}"
    return str(n)

def accent_spanish_number_words(text: str) -> str:
    replacements = {
        "dieciseis": "dieciséis",
        "veintidos": "veintidós",
        "veintitres": "veintitrés",
        "veintiseis": "veintiséis",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

def salary_words(value: str | None) -> str:
    amount = currency(value)
    whole = int(amount)
    cents = int(round((amount - whole) * 100))
    words = number_to_words_es(whole)
    words = accent_spanish_number_words(words)
    return f"{words.upper()} {cents:02d}/100"

MONTHS_ES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
]

def date_words_es(value: datetime | None = None) -> str:
    current = value or datetime.now()
    day_words = accent_spanish_number_words(number_to_words_es(current.day))
    year_words = accent_spanish_number_words(number_to_words_es(current.year))
    return f"{day_words} de {MONTHS_ES[current.month]} de {year_words}"

def date_long_words_es(value: datetime | None = None) -> str:
    current = value or datetime.now()
    day_words = accent_spanish_number_words(number_to_words_es(current.day))
    year_words = accent_spanish_number_words(number_to_words_es(current.year))
    return f"{day_words} días del mes de {MONTHS_ES[current.month]} del {year_words}"

def currency(value: str | None) -> float:
    if not value:
        return 0.0
    cleaned = re.sub(r"[^0-9.]", "", value.replace(",", ""))
    try:
        return float(cleaned)
    except ValueError:
        return 0.0

def calculate_isr(taxable_income: float) -> float:
    if taxable_income <= 550.00:
        return 0.0
    if taxable_income <= 895.24:
        return 17.67 + ((taxable_income - 550.00) * 0.10)
    if taxable_income <= 2038.10:
        return 60.00 + ((taxable_income - 895.24) * 0.20)
    return 288.57 + ((taxable_income - 2038.10) * 0.30)

def calculate_payroll_deductions(salary_value: str | None) -> dict[str, str]:
    gross = currency(salary_value)
    isss = min(gross, 1000.00) * 0.03
    afp = gross * 0.0725
    taxable = max(gross - isss - afp, 0)
    isr = calculate_isr(taxable)
    total = isss + afp + isr
    return {
        "isss": f"${isss:.2f}",
        "afp": f"${afp:.2f}",
        "isr": f"${isr:.2f}",
        "total_deductions": f"${total:.2f}",
        "net": f"${max(gross - total, 0):.2f}",
    }

def extract_pdf_text(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)

def find_first(patterns: list[str], text: str) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
        if match:
            return " ".join(match.group(1).split())
    return ""

def guess_employee_from_text(text: str) -> dict[str, str]:
    normalized = re.sub(r"[ \t]+", " ", text)
    dui = find_first([r"\b(\d{8}-\d)\b", r"DUI[:\s]+(\d{8}-\d)"], normalized)
    code = find_first([r"(?:codigo|cod\.?|empleado)[:\s#-]+([A-Z0-9-]{2,})"], normalized)
    salary = find_first([r"(?:salario|sueldo)[^\d$]{0,25}(\$?\s*\d+[,\d]*(?:\.\d{2})?)"], normalized)
    hire_date = find_first([r"(?:ingreso|fecha de ingreso)[:\s]+([0-9]{1,2}[/-][0-9]{1,2}[/-][0-9]{2,4})"], normalized)
    job = find_first([r"(?:puesto|cargo)[:\s]+([A-ZÁÉÍÓÚÑ0-9 /.-]{3,60})"], normalized)
    name = find_first(
        [
            r"(?:nombre|empleado)[:\s]+([A-ZÁÉÍÓÚÑ ]{5,80})",
            r"([A-ZÁÉÍÓÚÑ]{2,}(?: [A-ZÁÉÍÓÚÑ]{2,}){2,5})\s+(?:DUI|Documento|Cargo|Puesto)",
        ],
        normalized,
    )
    return {
        "dui": dui,
        "employee_code": code,
        "full_name": name,
        "job_title": job,
        "hire_date": hire_date,
        "end_date": "",
        "salary": salary,
        "salary_words": salary_words(salary) if salary else "",
        "isss": "",
        "afp": "",
        "isr": "",
        "personal_loan": "",
        "bank_loan": "",
        "fsv": "",
    }

def normalize_header(value: object) -> str:
    text = str(value or "").strip().lower()
    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ñ": "n", ".": "", "#": ""
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return re.sub(r"[^a-z0-9]+", "_", text).strip("_")

HEADER_ALIASES = {
    "dui": {"dui", "documento_unico_de_identidad", "documento", "identidad", "numero_doc_identidad", "num_doc_identidad", "no_doc_identidad"},
    "employee_code": {"codigo", "codigo_empleado", "cod_empleado", "empleado", "numero_empleado", "no_empleado"},
    "full_name": {"nombre", "nombre_completo", "empleado_nombre", "nombre_empleado", "trabajador"},
    "first_name": {"nombres", "primer_nombre"},
    "last_name": {"apellidos", "apellido", "primer_apellido"},
    "job_title": {"puesto", "cargo", "plaza", "posicion"},
    "hire_date": {"fecha_ingreso", "ingreso", "fecha_de_ingreso", "fecha_alta", "fecha_contratacion"},
    "end_date": {"fecha_fin", "fecha_retiro", "fecha_baja", "hasta"},
    "salary": {"salario", "sueldo", "salario_mensual", "sueldo_mensual", "salario_ordinario"},
    "salary_words": {"salario_letras", "sueldo_letras"},
    "isss": {"isss"},
    "afp": {"afp"},
    "isr": {"isr", "renta"},
    "personal_loan": {"prestamo_personal", "prestamos_personales"},
    "bank_loan": {"prestamo_bancario", "prestamos_bancarios"},
    "fsv": {"fsv"},
}

def cell_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.2f}"
    return str(value).strip()

def excel_currency(value: str) -> str:
    if not value:
        return ""
    amount = currency(value)
    return f"${amount:.2f}" if amount else value

def read_excel_rows(path: Path) -> list[tuple[object, ...]]:
    suffix = path.suffix.lower()
    if suffix == ".xls":
        book = xlrd.open_workbook(str(path))
        sheet = book.sheet_by_index(0)
        rows = []
        for row_idx in range(sheet.nrows):
            values = []
            for col_idx, cell in enumerate(sheet.row(row_idx)):
                val = cell.value
                if cell.ctype == xlrd.XL_CELL_DATE:
                    val = xlrd.xldate_as_datetime(val, book.datemode)
                values.append(val)
            rows.append(tuple(values))
        book.release_resources()
        return rows

    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    wb.close()
    return rows

def extract_employees_from_excel(path: Path) -> tuple[list[dict[str, str]], list[str]]:
    rows = read_excel_rows(path)
    if not rows:
        return [], []

    header_index = 0
    best_score = -1
    for idx, row in enumerate(rows[:10]):
        headers = [normalize_header(c) for c in row]
        score = sum(1 for h in headers for aliases in HEADER_ALIASES.values() if h in aliases)
        if score > best_score:
            best_score = score
            header_index = idx

    headers = [normalize_header(c) for c in rows[header_index]]
    column_map = {}
    for field, aliases in HEADER_ALIASES.items():
        for i, header in enumerate(headers):
            if header in aliases:
                column_map[field] = i
                break

    employees = []
    for row in rows[header_index + 1 :]:
        data = {field: cell_value(row[i]) if i < len(row) else "" for field, i in column_map.items()}
        if not data.get("full_name"):
            data["full_name"] = " ".join(part for part in [data.get("first_name", ""), data.get("last_name", "")] if part).strip()
        if not data.get("full_name") and not data.get("dui") and not data.get("employee_code"):
            continue
        if data.get("salary"):
            data["salary"] = excel_currency(data["salary"])
        if not data.get("salary_words") and data.get("salary"):
            data["salary_words"] = salary_words(data["salary"])
        employees.append(
            {
                "dui": data.get("dui", ""),
                "employee_code": data.get("employee_code", ""),
                "full_name": data.get("full_name", ""),
                "job_title": data.get("job_title", ""),
                "hire_date": data.get("hire_date", ""),
                "end_date": data.get("end_date", ""),
                "salary": data.get("salary", ""),
                "salary_words": data.get("salary_words", ""),
                "isss": data.get("isss", ""),
                "afp": data.get("afp", ""),
                "isr": data.get("isr", ""),
                "personal_loan": data.get("personal_loan", ""),
                "bank_loan": data.get("bank_loan", ""),
                "fsv": data.get("fsv", ""),
            }
        )
    return employees, sorted(column_map.keys())

def save_constancia_employee(data: dict[str, str]) -> str:
    nombre = data.get("full_name", "").strip()
    if not nombre:
        raise ValueError("El nombre del empleado es obligatorio.")
    
    dui = data.get("dui", "").strip()
    codigo = data.get("employee_code", "").strip()
    salario_raw = data.get("salary", "").strip()
    sal_val = currency(salario_raw)
    
    sal_letras = data.get("salary_words", "").strip()
    if not sal_letras and sal_val > 0:
        sal_letras = salary_words(str(sal_val))
        
    auto_deductions = calculate_payroll_deductions(str(sal_val))
    isss = data.get("isss", "").strip() or auto_deductions["isss"]
    afp = data.get("afp", "").strip() or auto_deductions["afp"]
    isr = data.get("isr", "").strip() or auto_deductions["isr"]
    
    personal = data.get("personal_loan", "").strip() or "$0.00"
    bank = data.get("bank_loan", "").strip() or "$0.00"
    fsv = data.get("fsv", "").strip() or "$0.00"
    
    cargo = data.get("job_title", "").strip()
    ingreso = data.get("hire_date", "").strip()
    fin = data.get("end_date", "").strip()
    
    db = get_db()
    existing = None
    if dui:
        existing = db.execute("SELECT codigo FROM empleados WHERE dui = ?", (dui,)).fetchone()
    if not existing and codigo:
        existing = db.execute("SELECT codigo FROM empleados WHERE codigo = ?", (codigo,)).fetchone()
        
    if existing:
        codigo = existing["codigo"]
        db.execute(
            """
            UPDATE empleados
            SET nombre=?, cargo=?, fecha_ingreso=?, salario_mensual=?,
                dui=?, fecha_fin=?, salario_letras=?, isss=?, afp=?, isr=?,
                prestamo_personal=?, prestamo_bancario=?, fsv=?
            WHERE codigo=?
            """,
            (nombre, cargo, ingreso, sal_val, dui, fin, sal_letras, isss, afp, isr, personal, bank, fsv, codigo)
        )
    else:
        if not codigo:
            row = db.execute("SELECT codigo FROM empleados ORDER BY CAST(codigo AS INTEGER) DESC LIMIT 1").fetchone()
            if row:
                try:
                    next_val = int(row["codigo"]) + 1
                    codigo = str(next_val).zfill(4)
                except ValueError:
                    codigo = "2000"
            else:
                codigo = "2000"
        
        db.execute(
            """
            INSERT INTO empleados (
                codigo, nombre, cargo, fecha_ingreso, salario_mensual, activo, dui,
                fecha_fin, salario_letras, isss, afp, isr, prestamo_personal, prestamo_bancario, fsv
            ) VALUES (?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (codigo, nombre, cargo, ingreso, sal_val, dui, fin, sal_letras, isss, afp, isr, personal, bank, fsv)
        )
    db.commit()
    db.close()
    return codigo

def get_active_template(doc_type: str) -> dict | None:
    db = get_db()
    row = db.execute(
        "SELECT * FROM templates WHERE doc_type = ? AND is_active = 1 ORDER BY id DESC LIMIT 1",
        (doc_type,),
    ).fetchone()
    db.close()
    return dict(row) if row else None

def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs) or paragraph.text
    new_text = full_text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

def render_docx_from_template(template_path: Path, employee: dict, doc_type: str) -> Path:
    doc = Document(template_path)
    generated_id = uuid.uuid4().hex[:10]
    out_path = Path(GENERATED_DIR) / f"{doc_type}-{employee['codigo']}-{generated_id}.docx"
    
    sal_val = employee.get("salario_mensual") or 0.0
    sal = f"${sal_val:.2f}"
    
    auto_deductions = calculate_payroll_deductions(str(sal_val))
    isss = employee.get("isss") or auto_deductions["isss"]
    afp = employee.get("afp") or auto_deductions["afp"]
    isr = employee.get("isr") or auto_deductions["isr"]
    personal = employee.get("prestamo_personal") or "$0.00"
    bank = employee.get("prestamo_bancario") or "$0.00"
    fsv = employee.get("fsv") or "$0.00"
    
    total_deducciones = sum(currency(x) for x in [isss, afp, isr, personal, bank, fsv])
    net = max(sal_val - total_deducciones, 0)
    fecha_emision = datetime.now()
    
    replacements = {
        "{{nombre_empleado}}": employee.get("nombre") or "",
        "{{dui}}": employee.get("dui") or "",
        "{{codigo_empleado}}": employee.get("codigo") or "",
        "{{puesto}}": employee.get("cargo") or "",
        "{{fecha_ingreso}}": employee.get("fecha_ingreso") or "",
        "{{fecha_fin}}": employee.get("fecha_fin") or "presente",
        "{{salario_numero}}": sal,
        "{{salario_letras}}": employee.get("salario_letras") or salary_words(str(sal_val)),
        "{{isss}}": isss,
        "{{afp}}": afp,
        "{{isr}}": isr,
        "{{prestamo_personal}}": personal,
        "{{prestamo_bancario}}": bank,
        "{{fsv}}": fsv,
        "{{total_deducciones}}": f"${total_deducciones:.2f}",
        "{{total_recibir}}": f"${net:.2f}",
        "{{fecha_emision}}": fecha_emision.strftime("%d/%m/%Y"),
        "{{fecha_emision_letras}}": date_words_es(fecha_emision),
        "{{ubicacion_adicional}}": f", específicamente en la ubicación de {employee.get('departamento')}" if employee.get("departamento") else "",
        
        # Mappings from original templates
        "DIEGO ALEXANDER HERNANDEZ MEBREÑO": employee.get("nombre") or "",
        "06373254-3": employee.get("dui") or "",
        "11 de Julio de 2024": employee.get("fecha_ingreso") or "",
        "CAPORAL": employee.get("cargo") or "",
        "CUATROCIENTOS CINCUENTA 00/100": employee.get("salario_letras") or salary_words(str(sal_val)),
        "$ 450.00": sal,
        "$450.00": sal,
        "$13.50": isss,
        "$32.62": afp,
        "$46.12": f"${total_deducciones:.2f}",
        "$403.88": f"${net:.2f}",
        "dos de junio de dos mil veintiséis": date_words_es(fecha_emision),
        "JUAN ALBERTO GONZÁLEZ SÁNCHEZ": employee.get("nombre") or "",
        "BAJA CERO": employee.get("cargo") or "",
        "14 de marzo del 2024": employee.get("fecha_ingreso") or "",
        "15 de mayo del 2026": employee.get("fecha_fin") or "presente",
        "veintidós días del mes de mayo del dos mil veintiséis": date_long_words_es(fecha_emision),
    }
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    doc.save(out_path)
    return out_path

def convert_to_pdf(docx_path: Path) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    candidate = Path(os.environ.get("PROGRAMFILES", "")) / "LibreOffice" / "program" / "soffice.exe"
    if not soffice and candidate.exists():
        soffice = str(candidate)
    if not soffice:
        return None
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=False,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    return pdf_path if pdf_path.exists() else None


# ===========================================================================
# ENDPOINTS DEL MÓDULO DE CONSTANCIAS
# ===========================================================================

@app.route("/admin/constancias")
@login_required
def admin_constancias():
    db = get_db()
    empleados = db.execute("SELECT codigo, nombre, cargo, dui FROM empleados WHERE activo=1 ORDER BY nombre").fetchall()
    templates = db.execute("SELECT * FROM templates ORDER BY created_at DESC").fetchall()
    docs = db.execute(
        """
        SELECT g.*, e.nombre, e.cargo FROM generated_documents g
        JOIN empleados e ON e.codigo = g.codigo_empleado
        ORDER BY g.created_at DESC
        """
    ).fetchall()
    db.close()
    return render_template(
        "admin_constancias.html",
        empleados=empleados,
        templates=templates,
        docs=docs
    )

@app.route("/admin/constancias/upload-excel", methods=["POST"])
@login_required
def admin_upload_excel():
    file = request.files.get("excel")
    if not file or file.filename == "":
        flash("Debes seleccionar un archivo Excel.", "error")
        return redirect(url_for("admin_constancias"))
    
    filename = secure_filename(file.filename)
    stored = os.path.join(UPLOAD_DIR, f"excel_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}")
    file.save(stored)
    
    try:
        employees, detected = extract_employees_from_excel(Path(stored))
        saved = 0
        skipped = 0
        for emp in employees:
            try:
                save_constancia_employee(emp)
                saved += 1
            except Exception as e:
                skipped += 1
        
        flash(f"Excel importado. Empleados actualizados/guardados: {saved}, omitidos: {skipped}.", "ok")
    except Exception as e:
        flash(f"Error al procesar el Excel: {e}", "error")
        
    return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/upload-report", methods=["POST"])
@login_required
def admin_upload_report():
    file = request.files.get("report")
    if not file or file.filename == "":
        flash("Debes seleccionar un archivo PDF.", "error")
        return redirect(url_for("admin_constancias"))
    
    filename = secure_filename(file.filename)
    stored = os.path.join(UPLOAD_DIR, f"pdf_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}")
    file.save(stored)
    
    try:
        text = extract_pdf_text(Path(stored))
        guessed = guess_employee_from_text(text)
        return render_template("admin_constancias_pdf_review.html", data=guessed, text_preview=text[:3000])
    except Exception as e:
        flash(f"Error al leer el PDF: {e}", "error")
        return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/save-employee", methods=["POST"])
@login_required
def admin_save_constancia_employee():
    data = {
        "full_name": request.form.get("full_name", ""),
        "dui": request.form.get("dui", ""),
        "employee_code": request.form.get("employee_code", ""),
        "job_title": request.form.get("job_title", ""),
        "hire_date": request.form.get("hire_date", ""),
        "end_date": request.form.get("end_date", ""),
        "salary": request.form.get("salary", ""),
        "salary_words": request.form.get("salary_words", ""),
        "isss": request.form.get("isss", ""),
        "afp": request.form.get("afp", ""),
        "isr": request.form.get("isr", ""),
        "personal_loan": request.form.get("personal_loan", ""),
        "bank_loan": request.form.get("bank_loan", ""),
        "fsv": request.form.get("fsv", "")
    }
    try:
        codigo = save_constancia_employee(data)
        flash(f"Empleado guardado correctamente (Código: {codigo}).", "ok")
    except Exception as e:
        flash(f"Error al guardar el empleado: {e}", "error")
    return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/upload-template", methods=["POST"])
@login_required
def admin_upload_template():
    doc_type = request.form.get("doc_type", "salario")
    file = request.files.get("template")
    if not file or file.filename == "":
        flash("Debes seleccionar un archivo .docx", "error")
        return redirect(url_for("admin_constancias"))
    
    filename = secure_filename(file.filename)
    folder = os.path.join(TEMPLATES_DIR, doc_type)
    os.makedirs(folder, exist_ok=True)
    stored = os.path.join(folder, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}")
    file.save(stored)
    
    db = get_db()
    db.execute("UPDATE templates SET is_active = 0 WHERE doc_type = ?", (doc_type,))
    db.execute(
        "INSERT INTO templates (doc_type, filename, stored_path, is_active, created_at) VALUES (?, ?, ?, 1, ?)",
        (doc_type, filename, stored, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    db.commit()
    db.close()
    
    flash(f"Nueva plantilla de {doc_type} cargada y activada.", "ok")
    return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/generate", methods=["POST"])
@login_required
def admin_generate_constancia():
    codigo = request.form.get("codigo")
    doc_type = request.form.get("doc_type", "salario")
    
    db = get_db()
    employee = db.execute("SELECT * FROM empleados WHERE codigo = ?", (codigo,)).fetchone()
    db.close()
    
    if not employee:
        flash("No se encontró el empleado seleccionado.", "error")
        return redirect(url_for("admin_constancias"))
        
    template = get_active_template(doc_type)
    if not template:
        flash(f"No hay una plantilla activa para constancias de {doc_type}. Súbela primero.", "error")
        return redirect(url_for("admin_constancias"))
        
    try:
        docx_path = render_docx_from_template(Path(template["stored_path"]), dict(employee), doc_type)
        pdf_path = convert_to_pdf(docx_path)
        
        db = get_db()
        db.execute(
            "INSERT INTO generated_documents (codigo_empleado, doc_type, docx_path, pdf_path, created_at, generado_por) VALUES (?, ?, ?, ?, ?, 'admin')",
            (codigo, doc_type, str(docx_path), str(pdf_path) if pdf_path else None, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        db.commit()
        db.close()
        
        flash(f"Constancia de {doc_type} generada correctamente.", "ok")
    except Exception as e:
        flash(f"Error al generar la constancia: {e}", "error")
        
    return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/delete/<int:id>", methods=["POST"])
@login_required
def admin_delete_constancia(id):
    db = get_db()
    row = db.execute("SELECT * FROM generated_documents WHERE id = ?", (id,)).fetchone()
    if not row:
        db.close()
        flash("La constancia no existe.", "error")
        return redirect(url_for("admin_constancias"))
        
    db.execute("DELETE FROM generated_documents WHERE id = ?", (id,))
    db.commit()
    db.close()
    
    for path_val in (row["docx_path"], row["pdf_path"]):
        if path_val:
            try:
                os.remove(path_val)
            except Exception:
                pass
                
    flash("Constancia eliminada correctamente.", "ok")
    return redirect(url_for("admin_constancias"))

@app.route("/admin/constancias/download/<path:filename>")
@login_required
def admin_download_constancia(filename):
    safe_name = os.path.basename(filename)
    path = os.path.join(GENERATED_DIR, safe_name)
    if not os.path.exists(path):
        abort(404)
    return send_from_directory(GENERATED_DIR, safe_name, as_attachment=True, download_name=safe_name)


# ===========================================================================
# AUTOSERVICIO DE CONSTANCIAS (PÚBLICO) & REPORTE DE VACACIONES (ADMIN)
# ===========================================================================

@app.route("/constancias")
def public_constancias():
    return render_template("solicitud_constancia.html")

@app.route("/api/constancia/generate", methods=["POST"])
def public_generate_constancia():
    dui = request.form.get("dui", "").strip()
    doc_type = request.form.get("doc_type", "salario")
    
    if not dui:
        flash("El número de DUI es obligatorio.", "error")
        return redirect(url_for("public_constancias"))
        
    db = get_db()
    employee = db.execute("SELECT * FROM empleados WHERE dui = ?", (dui,)).fetchone()
    db.close()
    
    if not employee:
        flash("No se encontró ningún empleado registrado con ese DUI.", "error")
        return redirect(url_for("public_constancias"))
        
    if not employee["activo"]:
        flash("El empleado asociado a este DUI no se encuentra activo.", "error")
        return redirect(url_for("public_constancias"))
        
    template = get_active_template(doc_type)
    if not template:
        flash(f"No hay una plantilla activa para constancias de {doc_type}.", "error")
        return redirect(url_for("public_constancias"))
        
    try:
        docx_path = render_docx_from_template(Path(template["stored_path"]), dict(employee), doc_type)
        pdf_path = convert_to_pdf(docx_path)
        
        db = get_db()
        db.execute(
            """
            INSERT INTO generated_documents (codigo_empleado, doc_type, docx_path, pdf_path, created_at, generado_por)
            VALUES (?, ?, ?, ?, ?, 'empleado')
            """,
            (employee["codigo"], doc_type, str(docx_path), str(pdf_path) if pdf_path else None, datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        db.commit()
        db.close()
        
        file_to_send = pdf_path if pdf_path else docx_path
        return send_from_directory(
            GENERATED_DIR, 
            file_to_send.name, 
            as_attachment=True, 
            download_name=file_to_send.name
        )
    except Exception as e:
        flash(f"Error al generar la constancia: {e}", "error")
        return redirect(url_for("public_constancias"))

@app.route("/admin/reporte-vacaciones")
@login_required
def admin_reporte_vacaciones():
    dep_filter = request.args.get("departamento", "").strip()
    month_filter = request.args.get("mes", "").strip()
    search_query = request.args.get("q", "").strip()
    
    db = get_db()
    departments = [row["departamento"] for row in db.execute(
        "SELECT DISTINCT departamento FROM empleados WHERE activo=1 AND departamento IS NOT NULL AND departamento != '' ORDER BY departamento"
    ).fetchall()]
    
    query = "SELECT * FROM empleados WHERE activo=1"
    params = []
    
    if dep_filter:
        query += " AND departamento = ?"
        params.append(dep_filter)
        
    if month_filter:
        m_val = month_filter.zfill(2)
        # Match using fecha_ingreso or fecha_contratacion month
        query += " AND (strftime('%m', fecha_ingreso) = ? OR strftime('%m', fecha_contratacion) = ?)"
        params.extend([m_val, m_val])
        
    if search_query:
        query += " AND (nombre LIKE ? OR codigo LIKE ? OR dui LIKE ?)"
        q_val = f"%{search_query}%"
        params.extend([q_val, q_val, q_val])
        
    query += " ORDER BY nombre"
    empleados_raw = db.execute(query, params).fetchall()
    db.close()
    
    empleados = []
    total_salarios = 0.0
    total_prima = 0.0
    total_dias_pendientes = 0.0
    total_valor_pendientes = 0.0
    total_costo_proyectado = 0.0
    
    for emp in empleados_raw:
        sal = emp["salario_mensual"] or 0.0
        dias_pend = emp["dias_vacacion_pendientes"] or 0.0
        
        prima = (sal / 30.0 * 15.0) * 0.30
        valor_dias = (sal / 30.0) * dias_pend
        total_emp = prima + valor_dias
        
        total_salarios += sal
        total_prima += prima
        total_dias_pendientes += dias_pend
        total_valor_pendientes += valor_dias
        total_costo_proyectado += total_emp
        
        month_num = None
        # Try to parse from fecha_ingreso
        if emp["fecha_ingreso"]:
            try:
                parts = emp["fecha_ingreso"].split("-")
                if len(parts) >= 2:
                    month_num = int(parts[1])
            except (ValueError, IndexError):
                pass
        
        month_name = MONTHS_ES[month_num].capitalize() if (month_num and 0 < month_num < len(MONTHS_ES)) else "N/A"
        
        empleados.append({
            "codigo": emp["codigo"],
            "nombre": emp["nombre"],
            "cargo": emp["cargo"],
            "departamento": emp["departamento"],
            "fecha_ingreso": emp["fecha_ingreso"],
            "fecha_contratacion": emp["fecha_contratacion"],
            "salario": sal,
            "dias_pendientes": dias_pend,
            "prima": prima,
            "valor_dias": valor_dias,
            "total": total_emp,
            "mes_aniversario": month_name
        })
        
    filtros = {
        "departamento": dep_filter,
        "mes": month_filter,
        "q": search_query
    }
    
    return render_template(
        "admin_reporte_vacaciones.html",
        empleados=empleados,
        departments=departments,
        filtros=filtros,
        total_salarios=total_salarios,
        total_prima=total_prima,
        total_dias_pendientes=total_dias_pendientes,
        total_valor_pendientes=total_valor_pendientes,
        total_costo_proyectado=total_costo_proyectado,
        months=MONTHS_ES
    )

@app.route("/admin/reporte-vacaciones/edit-days", methods=["POST"])
@login_required
def admin_reporte_vacaciones_edit_days():
    codigo = request.form.get("codigo")
    dias_s = request.form.get("dias_pendientes")
    
    if not codigo:
        return jsonify({"ok": False, "error": "El código del empleado es obligatorio."}), 400
        
    try:
        dias = float(dias_s) if dias_s else 0.0
    except ValueError:
        return jsonify({"ok": False, "error": "El valor de los días debe ser un número válido."}), 400
        
    db = get_db()
    cur = db.cursor()
    cur.execute("SELECT codigo FROM empleados WHERE codigo = ?", (codigo,))
    if not cur.fetchone():
        db.close()
        return jsonify({"ok": False, "error": "Empleado no encontrado."}), 404
        
    cur.execute("UPDATE empleados SET dias_vacacion_pendientes = ? WHERE codigo = ?", (dias, codigo))
    db.commit()
    db.close()
    
    return jsonify({"ok": True})

@app.route("/admin/reporte-vacaciones/export")
@login_required
def admin_reporte_vacaciones_export():
    dep_filter = request.args.get("departamento", "").strip()
    month_filter = request.args.get("mes", "").strip()
    search_query = request.args.get("q", "").strip()
    
    db = get_db()
    query = "SELECT * FROM empleados WHERE activo=1"
    params = []
    
    if dep_filter:
        query += " AND departamento = ?"
        params.append(dep_filter)
        
    if month_filter:
        m_val = month_filter.zfill(2)
        query += " AND (strftime('%m', fecha_ingreso) = ? OR strftime('%m', fecha_contratacion) = ?)"
        params.extend([m_val, m_val])
        
    if search_query:
        query += " AND (nombre LIKE ? OR codigo LIKE ? OR dui LIKE ?)"
        q_val = f"%{search_query}%"
        params.extend([q_val, q_val, q_val])
        
    query += " ORDER BY nombre"
    empleados_raw = db.execute(query, params).fetchall()
    db.close()
    
    output = io.StringIO()
    output.write(u'\ufeff')
    writer = csv.writer(output)
    
    writer.writerow([
        "Código", "Nombre", "Departamento", "Cargo", "Fecha Ingreso", "Fecha Contratación",
        "Salario Nominal ($)", "Prima Vacacional 30% ($)", "Días Vacación Pendientes", "Valor Días Pendientes ($)", "Total Proyección ($)"
    ])
    
    for emp in empleados_raw:
        sal = emp["salario_mensual"] or 0.0
        dias_pend = emp["dias_vacacion_pendientes"] or 0.0
        prima = (sal / 30.0 * 15.0) * 0.30
        valor_dias = (sal / 30.0) * dias_pend
        total_emp = prima + valor_dias
        
        writer.writerow([
            emp["codigo"],
            emp["nombre"],
            emp["departamento"] or "",
            emp["cargo"] or "",
            emp["fecha_ingreso"] or "",
            emp["fecha_contratacion"] or "",
            f"{sal:.2f}",
            f"{prima:.2f}",
            f"{dias_pend:.1f}",
            f"{valor_dias:.2f}",
            f"{total_emp:.2f}"
        ])
        
    filename = f"reporte_vacaciones_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    response = Response(output.getvalue(), mimetype="text/csv")
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    return response


@app.errorhandler(404)
def not_found(e):
    return render_template("404.html"), 404


if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=5000, debug=False)
else:
    init_db()
